"""Fixtures partagées pour les tests d'ingestion et de chunking."""

from pathlib import Path

import pytest


@pytest.fixture
def fixtures_dir() -> Path:
    """Répertoire contenant les documents réels fournis pour les tests d'intégration."""
    return Path(__file__).parent / "fixtures" / "documents"


@pytest.fixture
def simple_docx(tmp_path: Path) -> Path:
    """Génère un .docx minimal avec un titre, un paragraphe et un tableau."""
    from docx import Document

    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Ceci est un paragraphe de test décrivant le contexte du document.")
    doc.add_heading("Résultats financiers", level=1)
    doc.add_paragraph("Le chiffre d'affaires a augmenté de 12% sur l'exercice.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Trimestre"
    table.cell(0, 1).text = "Revenu"
    table.cell(1, 0).text = "Q1"
    table.cell(1, 1).text = "1200000"

    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def simple_pdf(tmp_path: Path) -> Path:
    """Génère un .pdf minimal avec un titre et un paragraphe."""
    from reportlab.pdfgen import canvas

    path = tmp_path / "sample.pdf"
    c = canvas.Canvas(str(path))
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 750, "Rapport Annuel")
    c.setFont("Helvetica", 11)
    c.drawString(72, 720, "Ce document presente les resultats financiers de l'exercice.")
    c.save()
    return path


@pytest.fixture
def corrupt_pdf(tmp_path: Path) -> Path:
    """Fichier avec extension .pdf mais contenu binaire invalide."""
    path = tmp_path / "corrupt.pdf"
    path.write_bytes(b"this is not a valid pdf content \x00\x01\x02")
    return path


@pytest.fixture
def corrupt_docx(tmp_path: Path) -> Path:
    """Fichier avec extension .docx mais contenu invalide (pas un zip OOXML)."""
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a real docx file")
    return path
